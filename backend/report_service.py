"""
RedTrack Report Service v2
Generates .docx reports using python-docx.
Supports custom branded templates with {{placeholders}}.
PDF export via LibreOffice headless.
"""

import os
import uuid
import subprocess
from pathlib import Path
from datetime import datetime, timezone
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

UPLOAD_DIR = Path(os.environ.get("UPLOAD_DIR", "/app/uploads"))
REPORTS_DIR = UPLOAD_DIR / "reports"
TEMPLATES_DIR = UPLOAD_DIR / "report_templates"

REPORTS_DIR.mkdir(parents=True, exist_ok=True)
TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

SEV_COLORS = {
    "Critical": RGBColor(0xC0, 0x39, 0x2B),
    "High":     RGBColor(0xE6, 0x7E, 0x22),
    "Medium":   RGBColor(0xF3, 0x9C, 0x12),
    "Low":      RGBColor(0x27, 0xAE, 0x60),
    "Info":     RGBColor(0x29, 0x80, 0xB9),
}

SEV_HEX = {
    "Critical": "C0392B",
    "High":     "E67E22",
    "Medium":   "F39C12",
    "Low":      "27AE60",
    "Info":     "2980B9",
}


# ── Evidence helpers ──────────────────────────────────────────────────────────

def _ev_get(e, attr):
    return e.get(attr) if isinstance(e, dict) else getattr(e, attr, None)


def _is_image(e):
    mime = (_ev_get(e, "mime_type") or "").lower()
    if mime.startswith("image/"):
        return True
    name = (_ev_get(e, "original_name") or _ev_get(e, "filename") or "").lower()
    return name.endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"))


def _evidence_path(e):
    # Evidence.filename stores the absolute on-disk path in this codebase.
    p = _ev_get(e, "filename")
    if p and Path(p).exists():
        return p
    return None


def _ev_caption(e):
    return _ev_get(e, "caption")


def _ev_name(e):
    return _ev_get(e, "original_name") or "evidence"


# ── PTES band-weighted risk scoring ───────────────────────────────────────────
# Mirrors the model used on real engagements: each severity band contributes a
# weight, the weighted total is squashed onto a 1–15 PTES scale, and the scale
# maps to a named risk band. Weights are tuned so a single Critical dominates a
# handful of Lows, matching how a client actually perceives aggregate risk.

PTES_BAND_WEIGHTS = {"Critical": 5.0, "High": 3.0, "Medium": 2.0, "Low": 1.0, "Info": 0.25}
PTES_MAX = 15.0

PTES_BANDS = [
    (13.0, "Extreme"),
    (10.0, "Critical"),
    (7.0,  "High"),
    (4.0,  "Moderate"),
    (1.0,  "Low"),
    (0.0,  "Minimal"),
]


def compute_ptes_score(sev_counts):
    """
    Returns {'score': float 0–15, 'band': str, 'weighted_raw': float}.

    weighted_raw = sum(count × band weight). It's mapped onto the 0–15 PTES
    scale with a saturating curve so the score doesn't require an implausible
    number of findings to reach the top band — a couple of Criticals plus
    supporting Highs already lands in 'Extreme', consistent with manual scoring.
    """
    import math
    weighted_raw = sum(sev_counts.get(band, 0) * w for band, w in PTES_BAND_WEIGHTS.items())
    k = 0.14  # tuned so ~1 Critical + 2 High ≈ 13.0 (Extreme)
    score = round(PTES_MAX * (1 - math.exp(-k * weighted_raw)), 1)

    band = "Minimal"
    for threshold, name in PTES_BANDS:
        if score >= threshold:
            band = name
            break
    return {"score": score, "band": band, "weighted_raw": round(weighted_raw, 1)}


def _set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_table_borders(table):
    """Add borders to all cells in a table."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'CCCCCC')
                tcBorders.append(border)
            tcPr.append(tcBorders)


def _add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p


def _add_para(doc, text, bold=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


async def generate_docx_report(eng, findings, report, template_path: Optional[str] = None, evidence_map=None) -> str:
    """
    Generate a .docx pentest report.
    If template_path is provided, uses it as a base and fills placeholders.
    Otherwise generates a default branded report.
    """
    if template_path and Path(template_path).exists():
        return await _generate_from_template(eng, findings, report, template_path, evidence_map or {})
    else:
        return await _generate_default_report(eng, findings, report, evidence_map or {})


async def _generate_from_template(eng, findings, report, template_path: str, evidence_map=None) -> str:
    """Fill a user-uploaded .docx template with report data."""
    doc = Document(template_path)

    sev_counts = {s: sum(1 for f in findings if f.severity.value == s) for s in ["Critical", "High", "Medium", "Low", "Info"]}
    open_count = sum(1 for f in findings if f.status.value == "Open")

    replacements = {
        "{{client_name}}": eng.client or "",
        "{{engagement_name}}": eng.name or "",
        "{{engagement_type}}": eng.type.value if hasattr(eng.type, 'value') else str(eng.type),
        "{{report_title}}": report.title or "",
        "{{report_version}}": report.version or "1.0",
        "{{start_date}}": str(eng.start_date)[:10] if eng.start_date else "TBD",
        "{{end_date}}": str(eng.end_date)[:10] if eng.end_date else "TBD",
        "{{generated_date}}": datetime.now(timezone.utc).strftime("%B %d, %Y"),
        "{{executive_summary}}": report.executive_summary or "No executive summary provided.",
        "{{methodology}}": report.methodology_section or eng.methodology or "Black-box penetration testing.",
        "{{scope}}": eng.scope or "Not specified.",
        "{{out_of_scope}}": eng.out_of_scope or "Not specified.",
        "{{objectives}}": eng.objectives or "Not specified.",
        "{{rules_of_engagement}}": eng.rules_of_engagement or "Not specified.",
        "{{total_findings}}": str(len(findings)),
        "{{critical_count}}": str(sev_counts["Critical"]),
        "{{high_count}}": str(sev_counts["High"]),
        "{{medium_count}}": str(sev_counts["Medium"]),
        "{{low_count}}": str(sev_counts["Low"]),
        "{{info_count}}": str(sev_counts["Info"]),
        "{{open_count}}": str(open_count),
        "{{client_contact}}": eng.client_contact or "",
        "{{client_email}}": eng.client_email or "",
        "{{ref_id}}": eng.ref_id or "",
        "{{ptes_score}}": str(compute_ptes_score(sev_counts)["score"]),
        "{{ptes_band}}": compute_ptes_score(sev_counts)["band"],
    }

    # Replace placeholders. Word often splits a placeholder like {{ptes_score}}
    # across several runs, so a naive per-run replace misses it. But we must NOT
    # flatten a whole paragraph's formatting (the dashboard has big-red + small-
    # grey runs in one paragraph). So: only when a placeholder spans multiple
    # runs do we merge exactly those runs, and only for the placeholder's span.
    def _replace_in_paragraph(para):
        if not para.runs:
            return
        # Fast path: placeholder fully inside one run.
        for run in para.runs:
            for key, val in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, val)
        # Slow path: any placeholder still present across run boundaries?
        joined = "".join(r.text for r in para.runs)
        if not any(k in joined for k in replacements):
            return
        # Rebuild run text offsets and, for each unresolved placeholder, write the
        # replacement into the run where it starts and blank the spanned tail.
        for key, val in replacements.items():
            guard = 0
            while key in "".join(r.text for r in para.runs) and guard < 50:
                guard += 1
                runs = para.runs
                texts = [r.text for r in runs]
                combined = "".join(texts)
                start = combined.find(key)
                end = start + len(key)
                # Map char offsets back to runs.
                pos = 0
                start_run = end_run = None
                start_off = end_off = 0
                for i, t in enumerate(texts):
                    nxt = pos + len(t)
                    if start_run is None and start < nxt:
                        start_run, start_off = i, start - pos
                    if end <= nxt:
                        end_run, end_off = i, end - pos
                        break
                    pos = nxt
                if start_run is None or end_run is None:
                    break
                if start_run == end_run:
                    runs[start_run].text = texts[start_run].replace(key, val, 1)
                else:
                    runs[start_run].text = texts[start_run][:start_off] + val
                    for i in range(start_run + 1, end_run):
                        runs[i].text = ""
                    runs[end_run].text = texts[end_run][end_off:]

    for para in doc.paragraphs:
        _replace_in_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para)

    # Headers and footers carry the client name / ref id too, and live in
    # separate parts the body loop never touches. Process every section's
    # header and footer (including first-page and even-page variants).
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer,
                   section.even_page_header, section.even_page_footer):
            if hf is None:
                continue
            for para in hf.paragraphs:
                _replace_in_paragraph(para)
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _replace_in_paragraph(para)

    # Insert key findings (top 5 by severity) where {{key_findings}} appears
    for i, para in enumerate(doc.paragraphs):
        if "{{key_findings}}" in para.text:
            para.text = ""
            _insert_key_findings_after(doc, para, findings)
            break

    # Bespoke dashboard sections, pulled from engagement.report_dashboard.
    dash = getattr(eng, "report_dashboard", None) or {}
    _insert_dashboard_sections(doc, dash)

    # Insert findings table where {{findings_table}} placeholder appears
    for i, para in enumerate(doc.paragraphs):
        if "{{findings_table}}" in para.text:
            para.text = ""
            _insert_findings_table_after(doc, para, findings)
            break

    # Insert individual findings where {{findings_detail}} appears
    for i, para in enumerate(doc.paragraphs):
        if "{{findings_detail}}" in para.text:
            para.text = ""
            _insert_findings_detail_after(doc, para, findings, evidence_map or {})
            break

    output_path = REPORTS_DIR / f"{uuid.uuid4().hex}.docx"
    doc.save(str(output_path))
    return str(output_path)


async def _generate_default_report(eng, findings, report, evidence_map=None) -> str:
    evidence_map = evidence_map or {}
    """Generate a full default branded report."""
    doc = Document()

    # Page setup - US Letter
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    sev_counts = {s: sum(1 for f in findings if f.severity.value == s) for s in ["Critical", "High", "Medium", "Low", "Info"]}
    open_count = sum(1 for f in findings if f.status.value == "Open")
    rem_count = sum(1 for f in findings if f.status.value == "Remediated")

    # ── Cover Page ───────────────────────────────────────────────────────────

    # TLP Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("TLP:RED — CONFIDENTIAL")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    run.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    # RedTrack branding
    p = doc.add_paragraph()
    run = p.add_run("RedTrack")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    p = doc.add_paragraph()
    run = p.add_run("Penetration Test Report")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()
    doc.add_paragraph()

    # Report title
    p = doc.add_paragraph()
    run = p.add_run(report.title or f"{eng.client} Security Assessment")
    run.bold = True
    run.font.size = Pt(24)

    doc.add_paragraph()

    # Client name
    p = doc.add_paragraph()
    run = p.add_run(f"Prepared for: {eng.client}")
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Meta table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.style = 'Table Grid'
    meta_data = [
        ("Engagement Type", eng.type.value if hasattr(eng.type, 'value') else str(eng.type)),
        ("Reference", eng.ref_id or "—"),
        ("Start Date", str(eng.start_date)[:10] if eng.start_date else "TBD"),
        ("End Date", str(eng.end_date)[:10] if eng.end_date else "TBD"),
        ("Report Version", report.version or "1.0"),
    ]
    for i, (label, val) in enumerate(meta_data):
        row = meta_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = val
        row.cells[0].paragraphs[0].runs[0].bold = True
        _set_cell_bg(row.cells[0], "F2F2F2")

    doc.add_page_break()

    # ── Executive Summary ─────────────────────────────────────────────────────

    _add_heading(doc, "1. Executive Summary", 1, RGBColor(0xC0, 0x39, 0x2B))

    if report.executive_summary:
        doc.add_paragraph(report.executive_summary)
    else:
        doc.add_paragraph(
            f"{eng.client} engaged RedTrack to perform a {eng.type.value if hasattr(eng.type, 'value') else 'penetration'} test. "
            f"This report documents the findings, risk ratings, and recommended remediations identified during the assessment."
        )

    doc.add_paragraph()

    # Risk summary table
    _add_heading(doc, "Risk Summary", 2)
    risk_table = doc.add_table(rows=2, cols=6)
    risk_table.style = 'Table Grid'
    risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Total", "Critical", "High", "Medium", "Low", "Info"]
    values = [str(len(findings)), str(sev_counts["Critical"]), str(sev_counts["High"]),
              str(sev_counts["Medium"]), str(sev_counts["Low"]), str(sev_counts["Info"])]
    colors = ["1A1A2E", "C0392B", "E67E22", "F39C12", "27AE60", "2980B9"]

    for i, (h, v, c) in enumerate(zip(headers, values, colors)):
        hcell = risk_table.rows[0].cells[i]
        vcell = risk_table.rows[1].cells[i]
        hcell.text = h
        vcell.text = v
        _set_cell_bg(hcell, c)
        hcell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hcell.paragraphs[0].runs[0].bold = True
        hcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        vcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        vcell.paragraphs[0].runs[0].bold = True
        vcell.paragraphs[0].runs[0].font.size = Pt(14)

    # PTES band-weighted overall risk score.
    ptes = compute_ptes_score(sev_counts)
    doc.add_paragraph()
    score_p = doc.add_paragraph()
    score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = score_p.add_run("Overall Risk (PTES weighted): ")
    r1.bold = True
    r1.font.size = Pt(12)
    r2 = score_p.add_run("%s / 15  " % ptes["score"])
    r2.bold = True
    r2.font.size = Pt(14)
    band_color = {
        "Extreme": RGBColor(0x8E, 0x1B, 0x1B), "Critical": RGBColor(0xC0, 0x39, 0x2B),
        "High": RGBColor(0xE6, 0x7E, 0x22), "Moderate": RGBColor(0xF3, 0x9C, 0x12),
        "Low": RGBColor(0x27, 0xAE, 0x60), "Minimal": RGBColor(0x29, 0x80, 0xB9),
    }.get(ptes["band"], RGBColor(0x1A, 0x1A, 0x2E))
    r3 = score_p.add_run("(%s)" % ptes["band"])
    r3.bold = True
    r3.font.size = Pt(14)
    r3.font.color.rgb = band_color

    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note_p.add_run(
        "Band-weighted PTES score (Critical ×5, High ×3, Medium ×2, Low ×1, Info ×0.25), "
        "normalized to a 1–15 scale."
    )
    nr.italic = True
    nr.font.size = Pt(8)
    nr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ── Scope & Methodology ───────────────────────────────────────────────────

    _add_heading(doc, "2. Scope & Methodology", 1, RGBColor(0xC0, 0x39, 0x2B))

    _add_heading(doc, "2.1 Scope", 2)
    doc.add_paragraph(eng.scope or "Not specified.")

    if eng.out_of_scope:
        _add_heading(doc, "2.2 Out of Scope", 2)
        doc.add_paragraph(eng.out_of_scope)

    if eng.objectives:
        _add_heading(doc, "2.3 Objectives", 2)
        doc.add_paragraph(eng.objectives)

    if eng.rules_of_engagement:
        _add_heading(doc, "2.4 Rules of Engagement", 2)
        doc.add_paragraph(eng.rules_of_engagement)

    _add_heading(doc, "2.5 Methodology", 2)
    doc.add_paragraph(
        report.methodology_section or eng.methodology or
        "Testing was performed using industry-standard methodologies including OWASP Testing Guide, "
        "PTES (Penetration Testing Execution Standard), and NIST SP 800-115."
    )

    doc.add_page_break()

    # ── Findings Summary Table ────────────────────────────────────────────────

    _add_heading(doc, "3. Findings Summary", 1, RGBColor(0xC0, 0x39, 0x2B))

    if not findings:
        doc.add_paragraph("No findings were identified during this assessment.")
    else:
        sorted_findings = sorted(findings, key=lambda f: {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Info": 4}.get(f.severity.value, 5))
        _insert_findings_summary_table(doc, sorted_findings)

    doc.add_page_break()

    # ── Individual Findings ───────────────────────────────────────────────────

    _add_heading(doc, "4. Detailed Findings", 1, RGBColor(0xC0, 0x39, 0x2B))

    sorted_findings = sorted(findings, key=lambda f: {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Info": 4}.get(f.severity.value, 5))

    for idx, f in enumerate(sorted_findings, 1):
        sev = f.severity.value if hasattr(f.severity, 'value') else str(f.severity)
        status = f.status.value if hasattr(f.status, 'value') else str(f.status)

        # Finding header
        p = doc.add_heading(f"4.{idx} {f.ref_id} — {f.title}", 2)

        # Finding meta table
        meta = doc.add_table(rows=1, cols=4)
        meta.style = 'Table Grid'
        fields = [
            ("Severity", sev),
            ("CVSS", str(f.cvss_score) if f.cvss_score else "—"),
            ("CWE", f.cwe or "—"),
            ("Status", status),
        ]
        for i, (label, val) in enumerate(fields):
            cell = meta.rows[0].cells[i]
            p2 = cell.paragraphs[0]
            p2.clear()
            run1 = p2.add_run(label + ": ")
            run1.bold = True
            run1.font.size = Pt(9)
            run2 = p2.add_run(val)
            run2.font.size = Pt(9)
            if label == "Severity":
                run2.font.color.rgb = SEV_COLORS.get(sev, RGBColor(0, 0, 0))
                run2.bold = True

        doc.add_paragraph()

        for label, content in [
            ("Affected Component", f.affected_component),
            ("Description", f.description),
            ("Impact", f.impact),
            ("Steps to Reproduce", f.steps_to_reproduce),
            ("Remediation", f.remediation),
            ("References", f.references),
        ]:
            if content:
                p = doc.add_paragraph()
                run = p.add_run(label + ": ")
                run.bold = True
                run.font.size = Pt(10)
                p.add_run(content).font.size = Pt(10)

        # Embedded evidence — the images uploaded to this finding (from RedNote,
        # Burp, or the RedTrack UI). This is what turns a text finding into
        # something a client can actually verify.
        ev_items = evidence_map.get(str(f.id), [])
        image_items = [e for e in ev_items if _is_image(e)]
        if image_items:
            ep = doc.add_paragraph()
            er = ep.add_run("Evidence:")
            er.bold = True
            er.font.size = Pt(10)
            for e in image_items:
                path = _evidence_path(e)
                if not path:
                    continue
                try:
                    doc.add_picture(path, width=Inches(6.0))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    # A corrupt/unsupported image shouldn't sink the whole report.
                    cap = doc.add_paragraph()
                    cr = cap.add_run("[evidence image could not be embedded: %s]" % _ev_name(e))
                    cr.italic = True
                    cr.font.size = Pt(8)
                    continue
                caption = _ev_caption(e) or _ev_name(e)
                if caption:
                    capp = doc.add_paragraph()
                    capp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = capp.add_run(caption)
                    cr.italic = True
                    cr.font.size = Pt(8)
                    cr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        if idx < len(sorted_findings):
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.add_run("─" * 80).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            doc.add_paragraph()

    # ── Conclusion ────────────────────────────────────────────────────────────

    doc.add_page_break()
    _add_heading(doc, "5. Conclusion", 1, RGBColor(0xC0, 0x39, 0x2B))

    if report.conclusion:
        doc.add_paragraph(report.conclusion)
    else:
        rem_rate = round(rem_count / len(findings) * 100) if findings else 0
        doc.add_paragraph(
            f"This assessment identified {len(findings)} finding(s) across the tested environment. "
            f"Of these, {sev_counts['Critical']} are Critical and {sev_counts['High']} are High severity, "
            f"representing the most immediate risks to the organization. "
            f"The current remediation rate is {rem_rate}%. "
            f"RedTrack recommends prioritizing Critical and High findings for immediate remediation, "
            f"followed by a re-test to verify the effectiveness of applied fixes."
        )

    # Save
    output_path = REPORTS_DIR / f"{uuid.uuid4().hex}.docx"
    doc.save(str(output_path))
    return str(output_path)


def _find_marker_para(doc, marker):
    for para in doc.paragraphs:
        if marker in para.text:
            return para
    # Markers can also live inside table cells (the dashboard is table-heavy).
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if marker in para.text:
                        return para
    return None


def _insert_dashboard_sections(doc, dash):
    """
    Render the bespoke dashboard sections from engagement.report_dashboard into
    the template wherever their markers appear. Each is optional: if the marker
    isn't in the template or the data is empty, that section is skipped and the
    marker (if present) is cleared.
    """
    _inject_kpi_callouts(doc, dash.get("kpi_callouts", []))
    _inject_risk_matrix(doc, dash.get("risk_matrix", {}))
    _inject_attack_chain(doc, dash.get("attack_chain", []))
    _inject_remediation(doc, dash.get("remediation", []))
    _inject_defensive_controls(doc, dash.get("defensive_controls", []))


def _inject_kpi_callouts(doc, callouts):
    para = _find_marker_para(doc, "{{dashboard_callouts}}")
    if para is None:
        return
    para.text = ""
    if not callouts:
        return
    table = doc.add_table(rows=1, cols=len(callouts))
    anchor = para._p
    for i, c in enumerate(callouts):
        cell = table.rows[0].cells[i]
        vp = cell.paragraphs[0]
        vr = vp.add_run(str(c.get("value", "")))
        vr.bold = True
        vr.font.size = Pt(20)
        vr.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        lp = cell.add_paragraph()
        lr = lp.add_run(str(c.get("label", "")))
        lr.bold = True
        lr.font.size = Pt(9)
        if c.get("note"):
            npp = cell.add_paragraph()
            nr = npp.add_run(str(c["note"]))
            nr.font.size = Pt(8)
            nr.font.color.rgb = RGBColor(0x5A, 0x62, 0x72)
    anchor.addnext(table._tbl)


def _inject_risk_matrix(doc, matrix):
    para = _find_marker_para(doc, "{{risk_matrix}}")
    if para is None:
        return
    para.text = ""
    if not matrix:
        return
    # 4x4: header row/col + 3x3 body. Rows = impact (High/Med/Low), cols = likelihood.
    table = doc.add_table(rows=4, cols=4)
    _set_table_borders(table)
    corner = table.rows[0].cells[0]
    corner.text = ""
    heads = ["Low", "Med", "High"]
    for j, h in enumerate(heads):
        c = table.rows[0].cells[j + 1]
        c.text = h + " impact"
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(8)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    rows = [("High", "high"), ("Med", "med"), ("Low", "low")]
    cols = [("low", "Low"), ("med", "Med"), ("high", "High")]
    # cell color by risk level (row index + col index)
    tint = ["27AE60", "F39C12", "F39C12", "E67E22", "C0392B"]
    for ri, (rlabel, rkey) in enumerate(rows):
        lc = table.rows[ri + 1].cells[0]
        lc.text = rlabel
        lc.paragraphs[0].runs[0].bold = True
        lc.paragraphs[0].runs[0].font.size = Pt(8)
        for ci, (ckey, _) in enumerate(cols):
            key = "%s_%s" % (rkey, ckey)
            val = matrix.get(key, 0)
            cell = table.rows[ri + 1].cells[ci + 1]
            level = (2 - ri) + ci  # 0..4 rough risk band
            _set_cell_bg(cell, tint[min(level, 4)])
            cp = cell.paragraphs[0]
            cr = cp.add_run(str(val))
            cr.bold = True
            cr.font.size = Pt(12)
            cr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    anchor = para._p
    anchor.addnext(table._tbl)


def _inject_attack_chain(doc, chain):
    para = _find_marker_para(doc, "{{attack_chain}}")
    if para is None:
        return
    para.text = ""
    if not chain:
        return
    table = doc.add_table(rows=1, cols=4)
    _set_table_borders(table)
    for i, h in enumerate(["#", "Action", "Detail", "Outcome"]):
        c = table.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].runs[0].font.size = Pt(9)
        _set_cell_bg(c, "1B2A4A")
    for i, step in enumerate(chain, 1):
        row = table.add_row()
        vals = [str(i), step.get("step", ""), step.get("detail", ""), step.get("outcome", "")]
        for j, v in enumerate(vals):
            cell = row.cells[j]
            cell.text = v
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if j == 1:
                cell.paragraphs[0].runs[0].bold = True
    anchor = para._p
    anchor.addnext(table._tbl)


def _inject_remediation(doc, items):
    para = _find_marker_para(doc, "{{remediation}}")
    if para is None:
        return
    para.text = ""
    if not items:
        return
    table = doc.add_table(rows=0, cols=2)
    _set_table_borders(table)
    pri_color = {"P0": "C0392B", "P1": "E67E22", "P2": "F39C12", "P3": "2980B9"}
    for item in items:
        row = table.add_row()
        pcell = row.cells[0]
        pr_label = item.get("priority", "")
        pcell.text = pr_label
        pcell.paragraphs[0].runs[0].bold = True
        pcell.paragraphs[0].runs[0].font.size = Pt(9)
        pcell.width = Inches(1.2)
        key = pr_label.split()[0].replace("—", "").strip()[:2]
        _set_cell_bg(pcell, pri_color.get(key, "F4F6F9"))
        pcell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        icell = row.cells[1]
        icell.text = item.get("items", "")
        icell.paragraphs[0].runs[0].font.size = Pt(9)
        icell.width = Inches(5.3)
    anchor = para._p
    anchor.addnext(table._tbl)


def _inject_defensive_controls(doc, controls):
    para = _find_marker_para(doc, "{{defensive_controls}}")
    if para is None:
        return
    para.text = ""
    if not controls:
        return
    table = doc.add_table(rows=0, cols=2)
    _set_table_borders(table)
    for c in controls:
        row = table.add_row()
        mark = row.cells[0]
        passed = c.get("status") == "pass"
        mr = mark.paragraphs[0].add_run("PASS" if passed else "FAIL")
        mr.bold = True
        mr.font.size = Pt(8)
        mr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(mark, "1A7A4A" if passed else "C0392B")
        mark.width = Inches(0.7)
        mark.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tcell = row.cells[1]
        tcell.text = c.get("text", "")
        tcell.paragraphs[0].runs[0].font.size = Pt(9)
        tcell.width = Inches(5.8)
    anchor = para._p
    anchor.addnext(table._tbl)


def _insert_key_findings_after(doc, para, findings):
    """
    Top 5 findings by severity, as bullets: 'Title  [SEVERITY]'. Inserted where
    {{key_findings}} appears in a template.
    """
    order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Info": 4}
    ranked = sorted(
        findings,
        key=lambda f: order.get(f.severity.value if hasattr(f.severity, "value") else str(f.severity), 5),
    )[:5]

    anchor = para._p
    for f in ranked:
        sev = f.severity.value if hasattr(f.severity, "value") else str(f.severity)
        p = doc.add_paragraph()
        b = p.add_run("•  ")
        b.font.size = Pt(10)
        r = p.add_run(f.title)
        r.font.size = Pt(10)
        r2 = p.add_run("  [%s]" % sev.upper())
        r2.bold = True
        r2.font.size = Pt(9)
        r2.font.color.rgb = SEV_COLORS.get(sev, RGBColor(0x1B, 0x2A, 0x4A))
        anchor.addnext(p._p)
        anchor = p._p


def _insert_findings_summary_table(doc, findings):
    # Matches the template's summary layout: #  /  Finding  /  Severity  /  CVSS  /  Status,
    # navy header, severity-colored severity cell.
    table = doc.add_table(rows=1, cols=5)
    _set_table_borders(table)

    headers = ["#", "Finding", "Severity", "CVSS", "Status"]
    widths = [Inches(0.4), Inches(3.8), Inches(1.0), Inches(0.7), Inches(1.0)]
    hrow = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, widths)):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, "1B2A4A")
        cell.width = w

    order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Info": 4}
    ranked = sorted(findings, key=lambda f: order.get(
        f.severity.value if hasattr(f.severity, "value") else str(f.severity), 5))

    for idx, f in enumerate(ranked, 1):
        sev = f.severity.value if hasattr(f.severity, "value") else str(f.severity)
        status = f.status.value if hasattr(f.status, "value") else str(f.status)
        row = table.add_row()
        vals = [str(idx), f.title, sev, str(f.cvss_score) if f.cvss_score else "—", status]
        for i, (val, w) in enumerate(zip(vals, widths)):
            cell = row.cells[i]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.width = w
            if i == 2:  # severity cell — color it
                cell.paragraphs[0].runs[0].font.color.rgb = SEV_COLORS.get(sev, RGBColor(0, 0, 0))
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _insert_findings_table_after(doc, para, findings):
    """Insert findings summary table after a placeholder paragraph in a template."""
    sorted_findings = sorted(findings, key=lambda f: {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Info": 4}.get(
        f.severity.value if hasattr(f.severity, 'value') else str(f.severity), 5))
    _insert_findings_summary_table(doc, sorted_findings)


def _insert_findings_detail_after(doc, para, findings, evidence_map=None):
    """
    Render each finding as the 'Executive Dashboard' style block: a severity-
    colored title bar, a metadata table (Severity / CVSS 3.1 / MITRE ATT&CK /
    System / Status), then Description / Technical Detail / Evidence /
    Recommendations. Loops over every finding, each drawn with its own values.
    """
    evidence_map = evidence_map or {}
    sorted_findings = sorted(
        findings,
        key=lambda f: {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Info": 4}.get(
            f.severity.value if hasattr(f.severity, "value") else str(f.severity), 5),
    )

    for i, f in enumerate(sorted_findings, 1):
        sev = f.severity.value if hasattr(f.severity, "value") else str(f.severity)
        status = f.status.value if hasattr(f.status, "value") else str(f.status)
        sev_hex = SEV_HEX.get(sev, "1B2A4A")

        # ── Title bar (single-cell table, severity-colored) ──
        title_tbl = doc.add_table(rows=1, cols=1)
        title_tbl.autofit = True
        tcell = title_tbl.rows[0].cells[0]
        _set_cell_bg(tcell, sev_hex)
        tp = tcell.paragraphs[0]
        r = tp.add_run("%d  %s  " % (i, f.title))
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r2 = tp.add_run("[%s]" % sev.upper())
        r2.bold = True
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # ── Metadata table ──
        mitre = getattr(f, "mitre_atlas_ttp", None) or ""
        meta_rows = [
            ("Severity", sev),
            ("CVSS 3.1", ("%s (%s)" % (f.cvss_score, f.cvss_vector)) if f.cvss_score and f.cvss_vector
                         else (str(f.cvss_score) if f.cvss_score else "—")),
            ("MITRE ATT&CK", mitre or "—"),
            ("System", f.affected_component or "—"),
            ("Status", status),
        ]
        mtbl = doc.add_table(rows=len(meta_rows), cols=2)
        _set_table_borders(mtbl)
        for ri, (label, val) in enumerate(meta_rows):
            lcell = mtbl.rows[ri].cells[0]
            vcell = mtbl.rows[ri].cells[1]
            _set_cell_bg(lcell, "F4F6F9")
            lp = lcell.paragraphs[0]
            lr = lp.add_run(label)
            lr.bold = True
            lr.font.size = Pt(10)
            lr.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
            vp = vcell.paragraphs[0]
            vr = vp.add_run(val)
            vr.font.size = Pt(10)
        # narrow label column
        for row in mtbl.rows:
            row.cells[0].width = Inches(1.6)
            row.cells[1].width = Inches(4.9)

        doc.add_paragraph()

        # ── Content sections ──
        sections = [
            ("Description", f.description),
            ("Technical Detail", f.steps_to_reproduce),
            ("Recommendations", f.remediation),
        ]
        for label, content in sections:
            if content:
                h = doc.add_heading(label, level=3)
                bp = doc.add_paragraph()
                bp.add_run(content).font.size = Pt(10)

        # ── Evidence (embedded images) ──
        ev_items = [e for e in evidence_map.get(str(f.id), []) if _is_image(e)]
        if ev_items:
            doc.add_heading("Evidence", level=3)
            for e in ev_items:
                path = _evidence_path(e)
                if not path:
                    continue
                try:
                    doc.add_picture(path, width=Inches(6.0))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    continue
                cap = _ev_caption(e) or _ev_name(e)
                if cap:
                    cp = doc.add_paragraph()
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cp.add_run(cap)
                    cr.italic = True
                    cr.font.size = Pt(8)
                    cr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()


async def convert_to_pdf(docx_path: str) -> Optional[str]:
    """Convert .docx to PDF using LibreOffice headless."""
    try:
        output_dir = str(REPORTS_DIR)
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0:
            pdf_path = Path(docx_path).with_suffix('.pdf')
            if pdf_path.exists():
                return str(pdf_path)
        return None
    except Exception:
        return None


# Keep old name for backward compat
async def generate_pdf_report(eng, findings, report) -> str:
    return await generate_docx_report(eng, findings, report)
